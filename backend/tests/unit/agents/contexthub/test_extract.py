import io

import httpx
import pytest
from docx import Document
from pypdf import PdfWriter
from respx import MockRouter

from agentdrops.agents.contexthub.extract import extract_file_text, fetch_url_text


def test_extract_file_text_txt() -> None:
    assert extract_file_text("txt", b"hello world") == "hello world"


def test_extract_file_text_csv() -> None:
    assert extract_file_text("csv", b"a,b\n1,2") == "a,b\n1,2"


def test_extract_file_text_docx() -> None:
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_file_text("docx", buffer.getvalue())

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_extract_file_text_pdf() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)

    # A blank page extracts to empty text — this only proves the pdf path runs without error
    # and returns a string, not that it finds text on a page with none.
    assert extract_file_text("pdf", buffer.getvalue()) == ""


def test_extract_file_text_unsupported_type_raises() -> None:
    with pytest.raises(ValueError, match="unsupported content_type"):
        extract_file_text("exe", b"whatever")


@pytest.mark.respx(base_url="https://internal.example.com")
async def test_fetch_url_text_strips_html_tags(respx_mock: MockRouter) -> None:
    respx_mock.get("/page").mock(
        return_value=httpx.Response(
            200, text="<html><body><script>ignored</script><p>Hello <b>world</b></p></body></html>"
        )
    )
    async with httpx.AsyncClient(base_url="https://internal.example.com") as client:
        text = await fetch_url_text("https://internal.example.com/page", client)

    assert "Hello" in text
    assert "world" in text
    assert "ignored" not in text
    assert "<p>" not in text
